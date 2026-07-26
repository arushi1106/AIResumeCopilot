from docx import Document
from docx.shared import Pt


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(14)


def generate_resume_docx(resume, output_path="generated_resume.docx"):
    """
    Generates an ATS-friendly DOCX resume from the tailored resume JSON.
    """

    doc = Document()

    # ===========================
    # Candidate Name
    # ===========================
    if resume.get("name"):
        title = doc.add_heading(resume["name"], level=0)
        title.style.font.size = Pt(22)

    # ===========================
    # Contact Details
    # ===========================
    contact = []

    if resume.get("email"):
        contact.append(resume["email"])

    if resume.get("phone"):
        contact.append(resume["phone"])

    if resume.get("linkedin"):
        contact.append(resume["linkedin"])

    if resume.get("github"):
        contact.append(resume["github"])

    if contact:
        doc.add_paragraph(" | ".join(contact))

    # ===========================
    # Professional Summary
    # ===========================
    if resume.get("professional_summary"):
        add_heading(doc, "Professional Summary")

        doc.add_paragraph(
            resume["professional_summary"]
        )

    # ===========================
    # Skills
    # ===========================
    if resume.get("skills"):
        add_heading(doc, "Skills")

        doc.add_paragraph(
            ", ".join(resume["skills"])
        )

    # ===========================
    # Experience
    # ===========================
    if resume.get("experience"):
        add_heading(doc, "Experience")

        for exp in resume["experience"]:

            company = exp.get("company", "")
            title = exp.get("title", "")

            doc.add_heading(
                f"{title} | {company}",
                level=2
            )

            for bullet in exp.get("bullets", []):
                doc.add_paragraph(
                    bullet,
                    style="List Bullet"
                )

    # ===========================
    # Projects
    # ===========================
    if resume.get("projects"):
        add_heading(doc, "Projects")

        for project in resume["projects"]:

            doc.add_heading(
                project.get("name", ""),
                level=2
            )

            doc.add_paragraph(
                project.get("description", "")
            )

    # ===========================
    # Education
    # ===========================
    if resume.get("education"):
        add_heading(doc, "Education")

        for edu in resume["education"]:

            institution = edu.get("institution", "")
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            date = edu.get("date", "")

            doc.add_heading(
                f"{degree} in {field}",
                level=2
            )

            doc.add_paragraph(institution)
            doc.add_paragraph(date)

            for detail in edu.get("details", []):
                doc.add_paragraph(
                    detail,
                    style="List Bullet"
                )

    # ===========================
    # Certifications
    # ===========================
    if resume.get("certifications"):
        add_heading(doc, "Certifications")

        for cert in resume["certifications"]:
            doc.add_paragraph(
                cert,
                style="List Bullet"
            )

    # ===========================
    # Awards
    # ===========================
    if resume.get("awards"):
        add_heading(doc, "Awards")

        for award in resume["awards"]:
            doc.add_paragraph(
                award,
                style="List Bullet"
            )

    # ===========================
    # Publications
    # ===========================
    if resume.get("publications"):
        add_heading(doc, "Publications")

        for publication in resume["publications"]:
            doc.add_paragraph(
                publication,
                style="List Bullet"
            )

    # ===========================
    # Languages
    # ===========================
    if resume.get("languages"):
        add_heading(doc, "Languages")

        doc.add_paragraph(
            ", ".join(resume["languages"])
        )

    # ===========================
    # Volunteering
    # ===========================
    if resume.get("volunteering"):
        add_heading(doc, "Volunteering")

        for item in resume["volunteering"]:
            doc.add_paragraph(
                item,
                style="List Bullet"
            )

    doc.save(output_path)

    return output_path